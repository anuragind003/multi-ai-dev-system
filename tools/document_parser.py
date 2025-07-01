import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any

# Set up logger for this module
logger = logging.getLogger(__name__)

# Use modern pypdf instead of PyPDF2
try:
    from pypdf import PdfReader
    PDF_LIBRARY = "pypdf"
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_LIBRARY = "PyPDF2"
        logger.warning("Using legacy PyPDF2. Consider upgrading to pypdf: pip install pypdf")
    except ImportError:
        PDF_LIBRARY = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentParser:
    """Enhanced document parser with modern libraries and comprehensive error handling."""
    
    def __init__(self):
        self.supported_formats = []
        
        if PDF_LIBRARY:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
        
        # Always support text files
        self.supported_formats.extend(['.txt', '.md', '.rst', '.rtf'])
        
        logger.info(f"Document Parser initialized with support for: {', '.join(self.supported_formats)}")
        if PDF_LIBRARY:
            logger.info(f"PDF support: {PDF_LIBRARY}")
    
    def parse(self, file_path: str) -> str:
        """Alias for parse_document to maintain API compatibility."""
        return self.parse_document(file_path)
        
    def parse_document(self, file_path: str) -> str:
        """Parse document with enhanced error handling and validation."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")
        
        # Check file size (reasonable limit: 50MB)
        file_size = file_path.stat().st_size
        if file_size > 50 * 1024 * 1024:
            raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f}MB. Maximum: 50MB")
        
        try:
            if file_ext == '.pdf':
                return self._parse_pdf_enhanced(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_docx_enhanced(file_path)
            elif file_ext in ['.txt', '.md', '.rst', '.rtf']:
                return self._parse_text_enhanced(file_path)
            else:
                raise ValueError(f"No parser implementation for {file_ext}")
                
        except Exception as e:
            raise RuntimeError(f"Failed to parse {file_path.name}: {str(e)}")
    
    def _parse_pdf_enhanced(self, file_path: Path) -> str:
        """Enhanced PDF parsing with better error handling and metadata extraction."""
        if not PDF_LIBRARY:
            raise ImportError("No PDF library available. Install pypdf: pip install pypdf")
        
        try:
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Extract metadata
                if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', 'Unknown'),
                        'author': pdf_reader.metadata.get('/Author', 'Unknown'),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from pages
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(f"--- Page {page_num} ---\n{page_text.strip()}")
                        else:
                            text_content.append(f"--- Page {page_num} ---\n[No extractable text]")
                    except Exception as e:
                        text_content.append(f"--- Page {page_num} (Error) ---\nFailed to extract: {str(e)}")
            
            if not any(content for content in text_content if '[No extractable text]' not in content and '(Error)' not in content):
                raise ValueError("PDF contains no extractable text. It may be image-based or corrupted.")
            
            # Add metadata header
            header = f"Document: {file_path.name}\n"
            if metadata:
                header += f"Title: {metadata.get('title', 'N/A')}\n"
                header += f"Author: {metadata.get('author', 'N/A')}\n"
                header += f"Pages: {metadata.get('pages', 'N/A')}\n"
            header += "=" * 50 + "\n\n"
            
            full_text = header + "\n\n".join(text_content)
            return full_text
            
        except Exception as e:
            raise RuntimeError(f"PDF parsing failed with {PDF_LIBRARY}: {str(e)}")
    
    def _parse_docx_enhanced(self, file_path: Path) -> str:
        """Enhanced DOCX parsing with table and header extraction."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Document metadata
            content_parts.append(f"Document: {file_path.name}")
            content_parts.append("=" * 50)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            if paragraphs:
                content_parts.append("CONTENT:")
                content_parts.extend(paragraphs)
            
            # Extract tables
            if doc.tables:
                content_parts.append("\nTABLES:")
                for i, table in enumerate(doc.tables, 1):
                    content_parts.append(f"\n--- Table {i} ---")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):  # Skip empty rows
                            content_parts.append(" | ".join(row_data))
            
            full_text = "\n".join(content_parts)
            
            if not full_text.strip():
                raise ValueError("DOCX appears to be empty")
            
            return full_text
            
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {str(e)}")
    
    def _parse_text_enhanced(self, file_path: Path) -> str:
        """Enhanced text parsing with encoding detection and validation."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    
                if content.strip():
                    # Add metadata header
                    header = f"Document: {file_path.name}\n"
                    header += f"Encoding: {encoding}\n"
                    header += f"Size: {len(content)} characters\n"
                    header += "=" * 50 + "\n\n"
                    
                    return header + content
                else:
                    raise ValueError("Text file appears to be empty")
                    
            except UnicodeDecodeError:
                continue  # Try next encoding
            except Exception as e:
                raise RuntimeError(f"Text file parsing failed: {str(e)}")
        
        raise RuntimeError(f"Could not decode text file with any encoding: {encodings}")
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Enhanced document metadata and statistics."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": "File not found", "supported": False}
        
        try:
            info = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_size_mb": round(file_path.stat().st_size / (1024 * 1024), 2),
                "file_extension": file_path.suffix.lower(),
                "supported": file_path.suffix.lower() in self.supported_formats,
                "pdf_library": PDF_LIBRARY if file_path.suffix.lower() == '.pdf' else None,
                "docx_available": DOCX_AVAILABLE if file_path.suffix.lower() in ['.docx', '.doc'] else None
            }
            
            if info["supported"]:
                try:
                    content = self.parse_document(str(file_path))
                    info.update({
                        "character_count": len(content),
                        "word_count": len(content.split()),
                        "line_count": content.count('\n') + 1,
                        "parse_successful": True
                    })
                except Exception as e:
                    info.update({
                        "parse_error": str(e),
                        "parse_successful": False
                    })
            
            return info
            
        except Exception as e:
            return {
                "file_name": file_path.name,
                "error": str(e),
                "supported": False
            }